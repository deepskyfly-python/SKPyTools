from docx import Document
from docx.shared import Inches

#filePath = 'demo.docx'
filePath = 'test.docx'
#获取文章全部内容
doc=Document(filePath)
#一级标题
print('read ==> Heading 1')
for p in doc.paragraphs:
    if p.style.name=='Heading 1':
        print(p.text)
 
 
#二级标题
print('read ==> Heading 2')
for p in doc.paragraphs:
    if p.style.name=='Heading 2':
        print(p.text)
 
 
#所有标题
import re
print('read ==> All Headings')
for p in doc.paragraphs:
    if re.match("^Heading \d+$",p.style.name):
        print(p.text)

#ListNumber
print('read ==> List Number')
for p in doc.paragraphs:
    if p.style.name=='List Number':
        print(p.text)        
 
#所有内容
print('read ==> Normal')
for p in doc.paragraphs:
    if p.style.name=='Normal':
        print(p.text)
#从前面可以看出，如果知道不同内容的style.name，那么要读这些内容是极其方便的，这些style.name可以通过：
#print(p.style.name)得到
 
for p in doc.paragraphs:
   if p.style.name=='级别3：黑体 13磅 20行距 段落前后20 左对齐':
        print(p.text)
        #输出对应内容

#所有内容
print('read ==> ALL')
for p in doc.paragraphs:
    print(p.style.name+" : " +p.text)